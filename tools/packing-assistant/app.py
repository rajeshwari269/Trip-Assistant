import os
import streamlit as st
import requests
from dotenv import load_dotenv
import google.generativeai as genai
from groq import Groq
import re
from datetime import datetime
from io import BytesIO
from docx import Document

load_dotenv()

# --- Config ---
st.set_page_config(page_title="Smart Packing Assistant", page_icon="🎒", layout="wide")

# Load API keys from environment
GEMINI_API_KEY   = os.getenv("GEMINI_API_KEY")
GROQ_API_KEY     = os.getenv("GROQ_API_KEY")
OPENCAGE_API_KEY = os.getenv("OPENCAGE_API_KEY")
WEATHER_API_KEY  = os.getenv("WEATHER_API_KEY")

missing = [k for k,v in {
    "GEMINI_API_KEY": GEMINI_API_KEY,
    "GROQ_API_KEY": GROQ_API_KEY,
    "OPENCAGE_API_KEY": OPENCAGE_API_KEY,
    "WEATHER_API_KEY": WEATHER_API_KEY,
}.items() if not v]
if missing:
    st.error(f"Missing environment variables: {', '.join(missing)}")
    st.stop()

genai.configure(api_key=GEMINI_API_KEY)

location_cache = {}

def get_lat_lon_from_opencage(location):
    if location in location_cache:
        return location_cache[location]
    try:
        url = f"https://api.opencagedata.com/geocode/v1/json?q={location}&key={OPENCAGE_API_KEY}"
        response = requests.get(url, timeout=20)
        response.raise_for_status()
        data = response.json()
        if data.get("results"):
            lat = data["results"][0]["geometry"]["lat"]
            lon = data["results"][0]["geometry"]["lng"]
            location_cache[location] = (lat, lon)
            return lat, lon
    except Exception:
        pass
    return None, None

def get_weather(lat, lon):
    try:
        url = f"https://api.openweathermap.org/data/2.5/forecast?lat={lat}&lon={lon}&appid={WEATHER_API_KEY}&units=metric"
        response = requests.get(url, timeout=20)
        response.raise_for_status()
        data = response.json()
        for forecast in data["list"]:
            if "weather" in forecast and "main" in forecast:
                description = forecast["weather"][0]["description"].capitalize()
                temperature = forecast["main"]["temp"]
                return f"{description}, {temperature}°C"
    except:
        pass
    return "Weather data not available"

def create_docx(weather, packing_list):
    buffer = BytesIO()
    doc = Document()
    doc.add_heading("Smart Packing Assistant", level=1)
    doc.add_paragraph(f"Weather Forecast: {weather}\n")
    doc.add_heading("Packing List:", level=2)
    for item in packing_list:
        doc.add_paragraph(f"- {item}")
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def extract_items_from_suggestions(suggestions):
    clean_list = []
    for line in suggestions.split("\n"):
        line = re.sub(r"[:•\-]+", "", line).strip()
        if line and not line.isdigit():
            clean_list.append(line)
    return clean_list

def generate_packing_list(location, activities, trip_type, trip_duration, people, minimalist=False):
    model = genai.GenerativeModel(model_name="gemini-2.0-flash-exp")
    chat_session = model.start_chat(history=[])

    people_info = "; ".join(
        [f"{p['name']} (Age: {p['age']}, Gender: {p['gender']}, Medical Needs: {p['medical_issues']})" for p in people]
    )
    activities = activities or "general activities"
    list_style = "minimalist packing list" if minimalist else "packing list"
    prompt = (
        f"Suggest a {list_style} for travelers going to {location} for {activities}. "
        f"The trip type is {trip_type}. Include only essential items based on the trip duration "
        f"of {trip_duration} days and travelers' details: {people_info}. "
        f"Make it concise, clear, and avoid repetition."
    )
    try:
        response = chat_session.send_message(prompt)
        return extract_items_from_suggestions(response.text)
    except:
        try:
            client = Groq(api_key=GROQ_API_KEY)
            completion = client.chat.completions.create(
                model="llama3-70b-8192",
                messages=[{"role": "user", "content": prompt}],
                temperature=1,
                max_tokens=512,
            )
            content = completion.choices[0].message.content
            return extract_items_from_suggestions(content)
        except:
            return []

def filter_excessive_items(packing_list, trip_duration):
    max_clothes = min(int(trip_duration) + 2, 10) if isinstance(trip_duration, int) else 30
    filtered_list = []
    for item in packing_list:
        if "TShirt" in item or "t-shirt" in item:
            item = re.sub(r"\d+", str(max_clothes), item)
        filtered_list.append(item)
    return filtered_list

# --- UI ---
st.title("🎒 Personalized Packing List Generator")
st.markdown("### Plan your trip smarter with weather-based personalized packing lists!")

st.sidebar.header("🌍 Trip Details")
location = st.sidebar.text_input("Enter a location:")
start_date = st.sidebar.date_input("Start Date", datetime.today())

trip_type = st.sidebar.selectbox("Select Trip Type", [
    "Business Trip", "Permanent Relocation", "Solo Trip",
    "Educational Trip", "Romantic Trip", "Adventure Trip"
])

end_date = None
if trip_type != "Permanent Relocation":
    end_date = st.sidebar.date_input("End Date", datetime.today())

activities = st.sidebar.text_area("Enter your activities (comma-separated):", value="")
list_type = st.sidebar.selectbox("Select Packing List Type", ["Detailed List", "Minimalist List"])
num_people = st.sidebar.number_input("Number of people:", min_value=1, step=1)

people = []
for i in range(num_people):
    with st.expander(f"👤 Person {i + 1} Details"):
        name = st.text_input(f"Name", key=f"name_{i}")
        gender = st.selectbox("Gender", ["Male", "Female", "Other"], key=f"gender_{i}")
        age = st.number_input("Age", min_value=0, step=1, key=f"age_{i}")
        medical_issues = st.text_area("Any medical issues", key=f"medical_{i}")
        people.append({"name": name, "gender": gender, "age": age, "medical_issues": medical_issues})

if st.sidebar.button("Generate Packing List 🧳"):
    if not location or not people:
        st.error("🚨 Please enter all required fields.")
    else:
        lat, lon = get_lat_lon_from_opencage(location)
        if lat and lon:
            trip_duration = 1
            if trip_type != "Permanent Relocation" and end_date:
                try:
                    trip_duration = max(1, (end_date - start_date).days)
                except:
                    pass
            else:
                trip_duration = "Permanent"

            weather = get_weather(lat, lon)
            packing_list = generate_packing_list(location, activities, trip_type, trip_duration, people, minimalist=(list_type == "Minimalist List"))
            packing_list = filter_excessive_items(packing_list, trip_duration)

            st.success("✅ Packing list generated successfully!")
            st.subheader("🌤 Weather Forecast")
            st.info(weather)
            st.subheader(f"📦 {list_type}")
            for item in packing_list:
                st.write(f"- {item}")

            docx_file = create_docx(weather, packing_list)
            st.download_button(
                "📅 Download Packing List",
                data=docx_file,
                file_name="Packing_List.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.error("Unable to retrieve location coordinates.")
